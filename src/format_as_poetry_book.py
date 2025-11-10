#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para formatar um arquivo .docx consolidado como livro de poemas.

Aplica formatação profissional para publicação de livro de poesias.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_number(section):
    """Adiciona numeração de páginas no rodapé."""
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adiciona campo de número de página
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    run.font.size = Pt(10)


def create_title_page(doc, title, author=""):
    """Cria página de título do livro."""
    # Espaço no topo
    for _ in range(8):
        doc.add_paragraph()
    
    # Título principal
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = 'Georgia'
    
    # Espaço
    for _ in range(3):
        doc.add_paragraph()
    
    # Subtítulo
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("Coletânea de Poemas")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.name = 'Georgia'
    
    # Espaço
    for _ in range(5):
        doc.add_paragraph()
    
    # Autor
    if author:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(author)
        run.font.size = Pt(16)
        run.font.name = 'Georgia'
    
    # Quebra de página
    doc.add_page_break()


def create_table_of_contents(doc, poem_titles):
    """Cria índice com lista de poemas."""
    # Título do índice
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("ÍNDICE")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'Georgia'
    
    doc.add_paragraph()
    
    # Lista de poemas
    for idx, title in enumerate(poem_titles, 1):
        entry = doc.add_paragraph()
        entry.paragraph_format.left_indent = Inches(0.5)
        run = entry.add_run(f"{title}")
        run.font.size = Pt(11)
        run.font.name = 'Georgia'
    
    # Quebra de página
    doc.add_page_break()


def format_poem_title(doc, title):
    """Formata o título do poema."""
    # Remove extensão .docx do título
    clean_title = title.replace('.docx', '').strip()
    
    # Adiciona espaço antes do título
    doc.add_paragraph()
    
    # Título do poema
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(clean_title)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Georgia'
    
    # Linha decorativa
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_para.add_run("• • •")
    run.font.size = Pt(10)
    run.font.name = 'Georgia'
    
    doc.add_paragraph()


def format_poem_content(doc, paragraphs):
    """Formata o conteúdo do poema."""
    for para in paragraphs:
        if para.text.strip():  # Ignora parágrafos vazios
            new_para = doc.add_paragraph()
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_para.paragraph_format.space_after = Pt(6)
            
            # Copia o texto com formatação
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.font.size = Pt(11)
                new_run.font.name = 'Georgia'
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
        else:
            # Parágrafo vazio = espaço entre estrofes
            doc.add_paragraph()


def format_as_poetry_book(input_file, output_file, book_title="Coletânea de Poemas", author=""):
    """
    Formata arquivo consolidado como livro de poemas.
    
    Args:
        input_file (str): Caminho do arquivo consolidado
        output_file (str): Caminho do arquivo formatado de saída
        book_title (str): Título do livro
        author (str): Nome do autor
        
    Returns:
        str: Caminho do arquivo gerado
    """
    print(f"Lendo arquivo: {input_file}")
    source_doc = Document(input_file)
    
    # Cria novo documento
    formatted_doc = Document()
    
    # Configura margens
    sections = formatted_doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    print("Extraindo títulos dos poemas...")
    
    # Extrai títulos dos poemas (parágrafos centralizados em negrito)
    poem_titles = []
    current_poem = []
    poems_content = []
    current_title = None
    
    for para in source_doc.paragraphs:
        # Detecta título de poema (centralizado e negrito)
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER and para.runs:
            is_title = any(run.bold and run.font.size and run.font.size.pt >= 12 for run in para.runs)
            
            if is_title and '.docx' in para.text:
                # Salva poema anterior
                if current_title and current_poem:
                    poems_content.append((current_title, current_poem))
                
                # Novo poema
                current_title = para.text.strip()
                poem_titles.append(current_title.replace('.docx', ''))
                current_poem = []
            else:
                # Conteúdo do poema
                if current_title:
                    current_poem.append(para)
        else:
            # Conteúdo do poema
            if current_title:
                current_poem.append(para)
    
    # Adiciona último poema
    if current_title and current_poem:
        poems_content.append((current_title, current_poem))
    
    print(f"Encontrados {len(poem_titles)} poemas")
    print("\nCriando livro formatado...")
    
    # Cria página de título
    print("  ✓ Página de título")
    create_title_page(formatted_doc, book_title, author)
    
    # Cria índice
    print("  ✓ Índice")
    create_table_of_contents(formatted_doc, [t.replace('.docx', '') for t in poem_titles])
    
    # Adiciona poemas formatados
    print("  ✓ Formatando poemas...")
    for idx, (title, content) in enumerate(poems_content, 1):
        print(f"    [{idx}/{len(poems_content)}] {title.replace('.docx', '')}")
        
        format_poem_title(formatted_doc, title)
        format_poem_content(formatted_doc, content)
        
        # Adiciona quebra de página (exceto no último poema)
        if idx < len(poems_content):
            formatted_doc.add_page_break()
    
    # Adiciona numeração de páginas
    print("  ✓ Numeração de páginas")
    add_page_number(formatted_doc.sections[0])
    
    # Salva documento
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    formatted_doc.save(output_file)
    
    print(f"\n{'='*60}")
    print("✓ LIVRO DE POEMAS CRIADO COM SUCESSO!")
    print(f"{'='*60}")
    print(f"Arquivo gerado: {output_file}")
    print(f"Total de poemas: {len(poem_titles)}")
    
    return str(output_file)


def main():
    """Função principal do script."""
    # Configurações
    INPUT_FILE = "output/consolidado.docx"
    OUTPUT_FILE = "output/livro_de_poemas.docx"
    BOOK_TITLE = "Coletânea de Poemas"
    AUTHOR = ""  # Deixe vazio ou adicione o nome do autor
    
    try:
        # Formata como livro de poemas
        output_file = format_as_poetry_book(
            input_file=INPUT_FILE,
            output_file=OUTPUT_FILE,
            book_title=BOOK_TITLE,
            author=AUTHOR
        )
        
    except FileNotFoundError as e:
        print(f"✗ Erro: {str(e)}")
        print("Certifique-se de que o arquivo de entrada existe.")
        
    except Exception as e:
        print(f"✗ Erro inesperado: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
