#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para consolidar múltiplos arquivos .docx em um único documento.

Lê todos os arquivos .docx de uma pasta e os consolida em um único arquivo,
preservando a formatação básica e inserindo quebras de página entre documentos.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_docx_files(input_folder):
    """
    Retorna uma lista ordenada alfabeticamente de arquivos .docx da pasta.
    
    Args:
        input_folder (str): Caminho da pasta de entrada
        
    Returns:
        list: Lista de caminhos de arquivos .docx ordenados alfabeticamente
    """
    input_path = Path(input_folder)
    
    if not input_path.exists():
        raise FileNotFoundError(f"Pasta não encontrada: {input_folder}")
    
    docx_files = sorted([f for f in input_path.glob("*.docx") if f.is_file()])
    
    return docx_files


def copy_paragraph(source_para, target_doc):
    """
    Copia um parágrafo de um documento para outro, preservando formatação básica.
    
    Args:
        source_para: Parágrafo de origem
        target_doc: Documento de destino
    """
    new_para = target_doc.add_paragraph()
    new_para.alignment = source_para.alignment
    new_para.style = source_para.style
    
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name


def add_document_title(target_doc, filename):
    """
    Adiciona o nome do arquivo como título centralizado no documento.
    
    Args:
        target_doc: Documento de destino
        filename (str): Nome do arquivo
    """
    title_para = target_doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = title_para.add_run(filename)
    run.bold = True
    run.font.size = Pt(14)
    
    # Adiciona linha em branco após o título
    target_doc.add_paragraph()


def consolidate_docx_files(input_folder, output_folder, output_filename="consolidado.docx", 
                          add_filename_titles=True):
    """
    Consolida múltiplos arquivos .docx em um único documento.
    
    Args:
        input_folder (str): Pasta com os arquivos .docx de entrada
        output_folder (str): Pasta onde o arquivo consolidado será salvo
        output_filename (str): Nome do arquivo de saída
        add_filename_titles (bool): Se True, adiciona o nome do arquivo como título
        
    Returns:
        str: Caminho do arquivo consolidado criado
    """
    # Obtém lista de arquivos .docx
    docx_files = get_docx_files(input_folder)
    
    if not docx_files:
        raise ValueError(f"Nenhum arquivo .docx encontrado em: {input_folder}")
    
    print(f"Encontrados {len(docx_files)} arquivos .docx para consolidar")
    
    # Cria documento consolidado
    consolidated_doc = Document()
    
    # Processa cada arquivo
    for idx, docx_file in enumerate(docx_files):
        print(f"Processando [{idx + 1}/{len(docx_files)}]: {docx_file.name}")
        
        try:
            # Lê o documento de origem
            source_doc = Document(docx_file)
            
            # Adiciona título com o nome do arquivo (opcional)
            if add_filename_titles:
                add_document_title(consolidated_doc, docx_file.name)
            
            # Copia todos os parágrafos
            for para in source_doc.paragraphs:
                copy_paragraph(para, consolidated_doc)
            
            # Adiciona quebra de página (exceto após o último documento)
            if idx < len(docx_files) - 1:
                consolidated_doc.add_page_break()
                
        except Exception as e:
            print(f"Erro ao processar {docx_file.name}: {str(e)}")
            print("Continuando com o próximo arquivo...")
            continue
    
    # Cria pasta de saída se não existir
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Salva documento consolidado
    output_file = output_path / output_filename
    consolidated_doc.save(output_file)
    
    print(f"\n✓ Arquivo consolidado criado com sucesso: {output_file}")
    print(f"✓ Total de documentos consolidados: {len(docx_files)}")
    
    return str(output_file)


def main():
    """Função principal do script."""
    # Configurações
    INPUT_FOLDER = "input"
    OUTPUT_FOLDER = "output"
    OUTPUT_FILENAME = "consolidado.docx"
    ADD_TITLES = True  # Adicionar nome dos arquivos como títulos
    
    try:
        # Executa consolidação
        output_file = consolidate_docx_files(
            input_folder=INPUT_FOLDER,
            output_folder=OUTPUT_FOLDER,
            output_filename=OUTPUT_FILENAME,
            add_filename_titles=ADD_TITLES
        )
        
        print(f"\n{'='*60}")
        print("CONSOLIDAÇÃO CONCLUÍDA COM SUCESSO!")
        print(f"{'='*60}")
        print(f"Arquivo gerado: {output_file}")
        
    except FileNotFoundError as e:
        print(f"✗ Erro: {str(e)}")
        print("Certifique-se de que a pasta de entrada existe.")
        
    except ValueError as e:
        print(f"✗ Erro: {str(e)}")
        
    except Exception as e:
        print(f"✗ Erro inesperado: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
